import json
import io
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models import (
    Document as DocumentModel,
    ArchitectureReview, ADRRecord, APISpec, GeneratedDoc
)
from backend.services import ai_service, audit_service

router = APIRouter(prefix="/api", tags=["generators"])


def doc_or_404(id: str, db: Session) -> DocumentModel:
    doc = db.query(DocumentModel).filter(DocumentModel.id == id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.post("/documents/{id}/generate-urs")
def generate_urs(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)

    def do_gen():
        return ai_service.generate_urs(doc.text)

    result = audit_service.with_audit(db, "generate_urs", {"document_id": id}, do_gen)

    gen = GeneratedDoc(
        document_id=id,
        content_type="urs",
        content=result.get("content", ""),
        needs_review=result.get("needs_review", False),
        confidence=result.get("confidence", "medium"),
    )
    db.add(gen)
    db.commit()
    db.refresh(gen)

    return {
        "id": gen.id,
        "created_at": gen.created_at.isoformat() + "Z",
        "document_id": gen.document_id,
        "content_type": gen.content_type,
        "content": gen.content,
        "needs_review": gen.needs_review,
        "confidence": gen.confidence,
    }


@router.post("/documents/{id}/generate-srs")
def generate_srs(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)

    def do_gen():
        return ai_service.generate_srs(doc.text)

    result = audit_service.with_audit(db, "generate_srs", {"document_id": id}, do_gen)

    gen = GeneratedDoc(
        document_id=id,
        content_type="srs",
        content=result.get("content", ""),
        needs_review=result.get("needs_review", False),
        confidence=result.get("confidence", "medium"),
    )
    db.add(gen)
    db.commit()
    db.refresh(gen)

    return {
        "id": gen.id,
        "created_at": gen.created_at.isoformat() + "Z",
        "document_id": gen.document_id,
        "content_type": gen.content_type,
        "content": gen.content,
        "needs_review": gen.needs_review,
        "confidence": gen.confidence,
    }


@router.post("/documents/{id}/generate-adr")
def generate_adr_endpoint(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)

    def do_gen():
        return ai_service.generate_adr(doc.text)

    result = audit_service.with_audit(db, "generate_adr", {"document_id": id}, do_gen)

    adr = ADRRecord(
        document_id=id,
        adr_json=json.dumps(result, ensure_ascii=False),
        needs_review=result.get("needs_review", False),
        confidence=result.get("confidence", "medium"),
    )
    db.add(adr)
    db.commit()
    db.refresh(adr)

    return adr_to_dict(adr, doc.title)


@router.post("/documents/{id}/design-api")
def design_api(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)

    def do_gen():
        return ai_service.generate_openapi(doc.text)

    result = audit_service.with_audit(db, "design_api", {"document_id": id}, do_gen)

    spec = APISpec(
        document_id=id,
        openapi_json=result.get("openapi_json", ""),
        openapi_yaml=result.get("openapi_yaml", ""),
    )
    db.add(spec)
    db.commit()
    db.refresh(spec)

    return {
        "id": spec.id,
        "created_at": spec.created_at.isoformat() + "Z",
        "document_id": spec.document_id,
        "openapi_json": spec.openapi_json,
        "openapi_yaml": spec.openapi_yaml,
        "document_title": doc.title,
    }


@router.post("/documents/{id}/recommend-architecture")
def recommend_architecture_endpoint(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)

    def do_gen():
        return ai_service.recommend_architecture(doc.text)

    result = audit_service.with_audit(
        db, "recommend_architecture", {"document_id": id}, do_gen
    )

    arch = ArchitectureReview(
        document_id=id,
        recommendation_json=json.dumps(result, ensure_ascii=False),
        needs_review=result.get("needs_review", False),
    )
    db.add(arch)
    db.commit()
    db.refresh(arch)

    return {
        "id": arch.id,
        "created_at": arch.created_at.isoformat() + "Z",
        "document_id": arch.document_id,
        "recommendation_json": result,
        "needs_review": arch.needs_review,
        "document_title": doc.title,
    }


@router.get("/documents/{id}/export/docx")
def export_docx(id: str, db: Session = Depends(get_db)):
    doc = doc_or_404(id, db)
    try:
        from docx import Document as DocxDocument
        d = DocxDocument()
        d.add_heading(doc.title, 0)
        d.add_paragraph(f"Type: {doc.doc_type}")
        d.add_paragraph(f"Created: {doc.created_at.isoformat()}")
        d.add_heading("Content", 1)
        d.add_paragraph(doc.text)

        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={doc.title[:50]}.docx"},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")


# ─── ADR List/Get ─────────────────────────────────────────────────────────────

router_adr = APIRouter(prefix="/api/adr", tags=["generators"])


def adr_to_dict(adr: ADRRecord, doc_title: Optional[str] = None) -> dict:
    try:
        adr_content = json.loads(adr.adr_json)
    except Exception:
        adr_content = {}
    return {
        "id": adr.id,
        "created_at": adr.created_at.isoformat() + "Z",
        "document_id": adr.document_id,
        "adr_json": adr_content,
        "needs_review": adr.needs_review,
        "confidence": adr.confidence,
        "document_title": doc_title,
    }


@router_adr.get("")
def list_adr_records(
    document_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    q = db.query(ADRRecord)
    if document_id:
        q = q.filter(ADRRecord.document_id == document_id)
    adrs = q.order_by(ADRRecord.created_at.desc()).all()
    result = []
    for adr in adrs:
        doc_title = None
        if adr.document_id:
            doc = db.query(DocumentModel).filter(DocumentModel.id == adr.document_id).first()
            doc_title = doc.title if doc else None
        adr_dict = adr_to_dict(adr, doc_title)
        if status:
            if adr_dict.get("adr_json", {}).get("status") != status:
                continue
        result.append(adr_dict)
    return result


@router_adr.get("/{id}")
def get_adr_record(id: str, db: Session = Depends(get_db)):
    adr = db.query(ADRRecord).filter(ADRRecord.id == id).first()
    if not adr:
        raise HTTPException(status_code=404, detail="ADR not found")
    doc_title = None
    if adr.document_id:
        doc = db.query(DocumentModel).filter(DocumentModel.id == adr.document_id).first()
        doc_title = doc.title if doc else None
    return adr_to_dict(adr, doc_title)


# ─── Architecture Reviews ─────────────────────────────────────────────────────

router_arch = APIRouter(prefix="/api/architecture-reviews", tags=["generators"])


@router_arch.get("")
def list_architecture_reviews(db: Session = Depends(get_db)):
    reviews = db.query(ArchitectureReview).order_by(
        ArchitectureReview.created_at.desc()
    ).all()
    result = []
    for ar in reviews:
        doc_title = None
        doc = db.query(DocumentModel).filter(DocumentModel.id == ar.document_id).first()
        if doc:
            doc_title = doc.title
        try:
            rec = json.loads(ar.recommendation_json)
        except Exception:
            rec = {}
        result.append({
            "id": ar.id,
            "created_at": ar.created_at.isoformat() + "Z",
            "document_id": ar.document_id,
            "recommendation_json": rec,
            "needs_review": ar.needs_review,
            "document_title": doc_title,
        })
    return result


# ─── API Specs ────────────────────────────────────────────────────────────────

router_specs = APIRouter(prefix="/api/api-specs", tags=["generators"])


@router_specs.get("")
def list_api_specs(db: Session = Depends(get_db)):
    specs = db.query(APISpec).order_by(APISpec.created_at.desc()).all()
    result = []
    for s in specs:
        doc_title = None
        doc = db.query(DocumentModel).filter(DocumentModel.id == s.document_id).first()
        if doc:
            doc_title = doc.title
        result.append({
            "id": s.id,
            "created_at": s.created_at.isoformat() + "Z",
            "document_id": s.document_id,
            "openapi_json": s.openapi_json,
            "openapi_yaml": s.openapi_yaml,
            "document_title": doc_title,
        })
    return result


@router_specs.get("/{id}")
def get_api_spec(id: str, db: Session = Depends(get_db)):
    spec = db.query(APISpec).filter(APISpec.id == id).first()
    if not spec:
        raise HTTPException(status_code=404, detail="API spec not found")
    doc_title = None
    doc = db.query(DocumentModel).filter(DocumentModel.id == spec.document_id).first()
    if doc:
        doc_title = doc.title
    return {
        "id": spec.id,
        "created_at": spec.created_at.isoformat() + "Z",
        "document_id": spec.document_id,
        "openapi_json": spec.openapi_json,
        "openapi_yaml": spec.openapi_yaml,
        "document_title": doc_title,
    }
